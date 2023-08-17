from gensim.summarization import summarize
from docx import Document
import glob
import os

def read_docx(file_path):
    doc = Document(file_path)
    result = [p.text for p in doc.paragraphs]
    return " ".join(result)

def save_summary_to_docx(summary, title, output_path):
    summary_doc = Document()
    summary_doc.add_paragraph(title)  # Add title if needed
    summary_doc.add_paragraph(summary)
    summary_doc.save(output_path)

def main():
    source_dir = r"C:\\test files\\Summary\\file"
    docx_files = glob.glob(os.path.join(source_dir, '*.docx'))

    for target_file in docx_files:
        text = read_docx(target_file)
        title = text.split('\n')[0]  # Assuming the title is the first line
        summary = summarize(text, ratio=0.3)  # Summarize to 30% of the original text
        save_summary_to_docx(summary, title, target_file.replace('.docx', '_Summary_en.docx'))

if __name__ == "__main__":
    main()
