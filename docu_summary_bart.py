# -*- coding: utf-8 -*-

from transformers import BartForConditionalGeneration, BartTokenizer
import docx2txt
import os
from docx import Document
from googletrans import Translator

# 소스 디렉터리와 대상 파일 정의
source_dir = r"C:\test files\Summary\file"
target_file = os.path.join(source_dir, "Original.docx")

# Word 파일 불러오기
text = docx2txt.process(target_file)

# BART 모델과 토크나이저 불러오기
model = BartForConditionalGeneration.from_pretrained("facebook/bart-large-cnn")
tokenizer = BartTokenizer.from_pretrained("facebook/bart-large-cnn")

# 입력 텍스트를 BART 입력 형식으로 변환
# 주의: BART 모델은 일반적으로 최대 1024 토큰의 입력을 처리할 수 있으나, 여기서는 16384로 설정하였습니다.
inputs = tokenizer([text], max_length=16384, return_tensors='pt', truncation=True)

# 입력 텍스트의 토큰 수 계산
num_tokens = len(tokenizer.tokenize(text))

# 요약 생성 (입력 텍스트의 최소 50%를 유지)
summary_ids = model.generate(inputs['input_ids'], num_beams=4, max_length=int(num_tokens*0.5), early_stopping=True)
summary_text = [tokenizer.decode(g, skip_special_tokens=True, clean_up_tokenization_spaces=False) for g in summary_ids]

# 요약 결과를 Word 파일에 저장
doc = Document()
doc.add_paragraph(summary_text[0])
doc.save(os.path.join(source_dir, 'Summary_en.docx'))

# 번역기 초기화
translator = Translator()

# 요약을 한국어로 번역
translated_summary = translator.translate(summary_text[0], dest='ko').text

# 번역된 요약을 Word 파일에 저장
doc = Document()
doc.add_paragraph(translated_summary)
doc.save(os.path.join(source_dir, 'Summary_ko.docx'))
