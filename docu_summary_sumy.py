from deepl import Translator

from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.text_rank import TextRankSummarizer
from docx import Document
import glob
import os

# deepl API key = f335352c-8859-089e-c921-077d69843ef6
deepl_api_key = 'f335352c-8859-089e-c921-077d69843ef6'  # Replace with the actual key from your comments or environment variables

def translate_to_korean(text, api_key):
    translator = Translator(api_key)
    translated_text = translator.translate_text(text, target_lang="KO")
    return translated_text.text  # Extract the text from the TextResult object

def save_translation_to_docx(translated_text, title, output_path):
    translated_doc = Document()
    translated_doc.add_paragraph(title)  # Add title
    translated_doc.add_paragraph(translated_text)  # Add translated text as a paragraph
    translated_doc.save(output_path)

def read_docx(file_path):
    doc = Document(file_path)
    title = doc.paragraphs[0].text
    result = [p.text for p in doc.paragraphs[1:]]  # Exclude the title from the content
    return title, " ".join(result)

def save_summary_to_docx(summary_sentences, title, output_path):
    summary_doc = Document()
    summary_doc.add_paragraph(title)  # Add title
    for sentence in summary_sentences:
        summary_doc.add_paragraph(str(sentence))  # Add each sentence as a separate paragraph
    summary_doc.save(output_path)

def main():
    source_dir = r"C:\\test files\\Summary\\file"
    docx_files = glob.glob(os.path.join(source_dir, '*.docx'))
    docx_files = [file for file in docx_files if not os.path.basename(file).startswith('~$')]

    for target_file in docx_files:
        title, text = read_docx(target_file)
        parser = PlaintextParser.from_string(text, Tokenizer("english"))
        summarizer = TextRankSummarizer()
        summary_sentences = summarizer(parser.document, 10)  # Summarize to 10 sentences
        save_summary_to_docx(summary_sentences, title, target_file.replace('.docx', '_Summary_en.docx'))

        summary_sentences = summarizer(parser.document, 10)  # Summarize to 10 sentences
        summary_text = " ".join([str(sentence) for sentence in summary_sentences])
        translated_summary = translate_to_korean(summary_text, deepl_api_key)
        save_summary_to_docx(summary_sentences, title, target_file.replace('.docx', '_Summary_en.docx'))
        save_translation_to_docx(translated_summary, title, target_file.replace('.docx', '_Summary_ko.docx'))


if __name__ == "__main__":
    main()
