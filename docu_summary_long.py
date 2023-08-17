import docx2txt
from docx import Document
from transformers import pipeline, LongformerTokenizer, LongformerForSeq2SeqLM
import os

# 소스 디렉터리와 대상 파일 정의
source_dir = r"C:\test files\Summary\file"
target_file = os.path.join(source_dir, "Original.docx")

# Word 파일 불러오기
text = docx2txt.process(target_file)

# Longformer 모델과 토크나이저 불러오기
tokenizer = LongformerTokenizer.from_pretrained("allenai/longformer-base-4096")
model = LongformerForSeq2SeqLM.from_pretrained("allenai/longformer-base-4096")

# 입력 텍스트를 Longformer 입력 형식으로 변환
inputs = tokenizer.encode("summarize: " + text, return_tensors='pt', max_length=4096, truncation=True)

# 요약 생성
outputs = model.generate(inputs, max_length=int(len(text) * 0.2), min_length=40, length_penalty=2.0, num_beams=4, early_stopping=True)
summary_text = tokenizer.decode(outputs[0], skip_special_tokens=True)

# 한국어로 번역
translator = pipeline("translation_en_to_ko", model="Helsinki-NLP/opus-mt-en-ko", tokenizer="Helsinki-NLP/opus-mt-en-ko")
translation = translator(summary_text)
translated_summary_text = translation[0]['translation_text']

# 요약 결과를 원본 언어와 한국어로 각각 Word 파일에 저장
doc = Document()
doc.add_paragraph(summary_text)
doc.save(os.path.join(source_dir, 'Summary_EN.docx'))

doc = Document()
doc.add_paragraph(translated_summary_text)
doc.save(os.path.join(source_dir, 'Summary_KO.docx'))
