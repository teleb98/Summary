from sumy.summarizers.lex_rank import LexRankSummarizer
from sumy.nlp.tokenizers import Tokenizer
from sumy.parsers.plaintext import PlaintextParser
from deepl import Translator
from docx import Document
import glob
import os
import shutil 

from sumy.summarizers.lsa import LsaSummarizer

from nltk.tokenize import sent_tokenize
import nltk
nltk.download('punkt')  # Download the required package for sentence tokenization


deepl_api_key = 'f335352c-8859-089e-c921-077d69843ef6'

def translate_to_korean(text, api_key):
    translator = Translator(api_key)
    translated_text = translator.translate_text(text, target_lang="KO")
    return translated_text.text

def read_docx(file_path):
    doc = Document(file_path)
    title = doc.paragraphs[0].text
    result = [p.text for p in doc.paragraphs[1:]]
    return title, " ".join(result)

def save_summary_to_docx(summary_sentences, title, output_path):
    summary_doc = Document()
    summary_doc.add_paragraph(title)
    for sentence in summary_sentences:
        summary_doc.add_paragraph(str(sentence))
    summary_doc.save(output_path)

def save_translation_to_docx(translated_text, title, output_path):
    translated_doc = Document()
    translated_doc.add_paragraph(title)
    translated_doc.add_paragraph(translated_text)
    translated_doc.save(output_path)

def save_original_to_destination(original_file, destination_path):
    destination_file = os.path.join(destination_path, os.path.basename(original_file))
    shutil.copy(original_file, destination_file)

def main():
    source_dir = r"C:\\test files\\Summary\\file"
    destination_dir = r"C:\\test files\\Summary\\translation"

    docx_files = glob.glob(os.path.join(source_dir, '*.docx'))
    docx_files = [file for file in docx_files if not os.path.basename(file).startswith('~$')]

    for target_file in docx_files:
        # Save the original file to destination
        save_original_to_destination(target_file, destination_dir)
        
        title, text = read_docx(target_file)
        original_sentences = sent_tokenize(text)
        original_sentence_count = len(original_sentences)

        print(f"Original sentence count for {target_file}: {original_sentence_count}")

        # Translate the original text to Korean (Cost saviing)
        # translated_original_text = translate_to_korean(text, deepl_api_key)
        # save_translation_to_docx(translated_original_text, title, os.path.join(destination_dir, os.path.basename(target_file.replace('.docx', '_Original_ko.docx'))))

        # Calculate the number of sentences for the summary (33% of original sentences)
        summary_sentence_count = len(original_sentences) // 3

        parser = PlaintextParser.from_string(text, Tokenizer("english"))
        summarizer = LsaSummarizer()  # Use LSA
        summary_sentences = summarizer(parser.document, summary_sentence_count)

        print(f"Summary sentence count for {target_file}: {summary_sentence_count}")

        # Save the English summary
        save_summary_to_docx(summary_sentences, title, os.path.join(destination_dir, os.path.basename(target_file.replace('.docx', '_Summary_en.docx'))))

        # Convert the summary to individual sentences for translation
        summary_text = " ".join([str(sentence) for sentence in summary_sentences])
        translated_summary = translate_to_korean(summary_text, deepl_api_key)
        
        # Save translated and summary files to destination
        save_translation_to_docx(translated_summary, title, os.path.join(destination_dir, os.path.basename(target_file.replace('.docx', '_Summary_ko.docx'))))

if __name__ == "__main__":
    main()
