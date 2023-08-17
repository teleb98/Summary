import glob
import os
from collections import defaultdict
from docx import Document
from docx.shared import Pt
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
from pkg_resources import parse_version
import re
import sklearn
from sklearn.feature_extraction.text import TfidfVectorizer
import nltk

nltk.download('punkt')
nltk.download('stopwords')


def extract_urls(text):
    urls = re.findall('https?://(?:[-\w.]|(?:%[\da-fA-F]{2}))+', text)
    return urls


def read_docx(file_path):
    doc = Document(file_path)
    result = [p.text for p in doc.paragraphs]
    title = result[0]
    title_font_size = doc.paragraphs[0].runs[0].font.size
    return " ".join(result), len(result), title, title_font_size


def summarize_text(text, num_paragraphs):
    summary_ratio = 0.5 if num_paragraphs < 10 else (0.4 if num_paragraphs < 20 else 0.3)
    sentences = sent_tokenize(text)
    vectorizer = TfidfVectorizer(stop_words='english')
    tfidf_matrix = vectorizer.fit_transform(sentences)
    word2tfidf = dict(zip(vectorizer.get_feature_names_out() if parse_version(sklearn.__version__) >= parse_version('0.24') else vectorizer.get_feature_names(), vectorizer.idf_))
    sentence_scores = defaultdict(int)
    for i, sentence in enumerate(sentences):
        for word in word_tokenize(sentence.lower()):
            if word in word2tfidf:
                sentence_scores[i] += word2tfidf[word]
    num_summary_sentences = int(len(sentences) * summary_ratio)
    top_sentence_indices = sorted(sentence_scores, key=sentence_scores.get, reverse=True)[:num_summary_sentences]
    top_sentence_indices.sort()
    summary_sentences = [sentences[i] for i in top_sentence_indices]
    return summary_sentences, top_sentence_indices


def add_title_to_doc(title, title_font_size, doc):
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run(title)
    title_run.font.size = title_font_size


def add_summary_to_doc(summary_sentences, doc):
    for sentence in summary_sentences:
        doc.add_paragraph(sentence)


def copy_last_paragraph_url(original_doc_path, summarized_text):
    doc = Document(original_doc_path)
    last_paragraph = doc.paragraphs[-1].text
    urls = extract_urls(last_paragraph)
    if urls:
        summarized_text += "\\n\\nSource: " + urls[0]
    return summarized_text


def main():
    source_dir = r"C:\\test files\\Summary\\file"
    docx_files = glob.glob(os.path.join(source_dir, '*.docx'))

    for target_file in docx_files:
        text, num_paragraphs, title, title_font_size = read_docx(target_file)
        summary_sentences, _ = summarize_text(text, num_paragraphs)
        summary_doc = Document()
        add_title_to_doc(title, title_font_size, summary_doc)
        add_summary_to_doc(summary_sentences, summary_doc)
        summarized_text = copy_last_paragraph_url(target_file, summary_doc.paragraphs[-1].text)
        summary_doc.paragraphs[-1].text = summarized_text
        summary_doc.save(target_file.replace('.docx', '_Summary_en.docx'))


if __name__ == "__main__":
    main()
