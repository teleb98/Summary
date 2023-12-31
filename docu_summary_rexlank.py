from sumy.summarizers.lex_rank import LexRankSummarizer
from sumy.nlp.tokenizers import Tokenizer
from sumy.parsers.plaintext import PlaintextParser
from deepl import Translator
from docx import Document
import glob
import os

deepl_api_key = 'f335352c-8859-089e-c921-077d69843ef6'

def translate_to_korean(text, api_key):
    translator = Translator(api_key)
    translated_text = translator.translate_text(text, target_lang="KO")
    return translated_text.text

def read_docx(file_path):
    doc = Document(file_path)
    title = doc.paragraphs[0].text
    result = [p.text for p in doc.paragraphs[1:]]
    return title, " ".join(result)

def save_summary_to_docx(summary_sentences, title, output_path):
    summary_doc = Document()
    summary_doc.add_paragraph(title)
    for sentence in summary_sentences:
        summary_doc.add_paragraph(str(sentence))
    summary_doc.save(output_path)

def save_translation_to_docx(translated_text, title, output_path):
    translated_doc = Document()
    translated_doc.add_paragraph(title)
    translated_doc.add_paragraph(translated_text)
    translated_doc.save(output_path)

def main():
    source_dir = r"C:\\test files\\Summary\\file"
    docx_files = glob.glob(os.path.join(source_dir, '*.docx'))
    docx_files = [file for file in docx_files if not os.path.basename(file).startswith('~$')]

    for target_file in docx_files:
        title, text = read_docx(target_file)
        parser = PlaintextParser.from_string(text, Tokenizer("english"))
        summarizer = LexRankSummarizer()
        summary_sentences = summarizer(parser.document, 10)

        # Save the English summary
        save_summary_to_docx(summary_sentences, title, target_file.replace('.docx', '_Summary_en.docx'))

        # Convert the summary to individual sentences for translation
        summary_text = " ".join([str(sentence) for sentence in summary_sentences])
        translated_summary = translate_to_korean(summary_text, deepl_api_key)
        
        # Save the Korean translation of the summary
        save_translation_to_docx(translated_summary, title, target_file.replace('.docx', '_Summary_ko.docx'))

if __name__ == "__main__":
    main()
